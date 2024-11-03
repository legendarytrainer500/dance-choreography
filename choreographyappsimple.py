import PySimpleGUI as sg
from docx import Document


class DanceApp:
    def __init__(self):
        # Data storage
        # Define a custom theme with your colors
        # Add your new theme colors and settings
        custom_theme = {'BACKGROUND': '#F0C1E1',
                        'TEXT': '#DA0C81',
                        'INPUT': '#c7e78b',
                        'TEXT_INPUT': '#000000',
                        'SCROLL': '#c7e78b',
                        'BUTTON': ('white', '#3D5300'),
                        'PROGRESS': ('#01826B', '#D0D0D0'),
                        'BORDER': 1,
                        'SLIDER_DEPTH': 0,
                        'PROGRESS_DEPTH': 0}

        # Register the custom theme
        sg.theme_add_new('CustomTheme', custom_theme)

        # Set the custom theme as the current theme
        sg.theme('CustomTheme')
        self.dance_moves = []
        self.lyrics = []
        self.lyric_moves = []  # List to store tuples of (lyric, [moves])
        self.current_selected_move = None

        # Define the layout for each tab
        self.layout = [
            [sg.TabGroup([
                [sg.Tab('Add Dance Moves', self.add_move_tab()),
                 sg.Tab('Lyrics Input', self.lyrics_tab()),
                 sg.Tab('Choreograph Lyrics', self.choreograph_tab()),
                 sg.Tab('Save Choreography', self.save_tab())]
            ])]
        ]

        # Create the window
        self.window = sg.Window("Dance Choreography App", self.layout, finalize=True)

    def add_move_tab(self):
        """Layout for the Add Dance Moves tab."""
        return [
            [sg.Text("Paste Dance Moves (one move per line):")],
            [sg.Multiline(size=(40, 10), key='-MOVE_INPUT-')],
            [sg.Button("Add Moves", key='-ADD_MOVES-')],
            [sg.Text("Current Moves Library:")],
            [sg.Listbox(values=self.dance_moves, size=(40, 10), key='-MOVE_LIST-')]
        ]

    def lyrics_tab(self):
        """Layout for the Lyrics Input tab."""
        return [
            [sg.Text("Paste Lyrics (one line per lyric):")],
            [sg.Multiline(size=(40, 15), key='-LYRICS_INPUT-')],
            [sg.Button("Submit Lyrics", key='-SUBMIT_LYRICS-')]
        ]

    def choreograph_tab(self):
        """Layout for the Choreograph Lyrics tab."""
        return [
            [sg.Text("Select a dance move and then click a lyric to assign it.")],
            [sg.Column([[sg.Text("Lyrics:")]], size=(300, 20)),
             sg.Column([[sg.Text("Dance Moves:")]], size=(300, 20))],
            [sg.Listbox(values=self.lyrics, size=(30, 15), key='-LYRICS_LIST-', enable_events=True),
             sg.Listbox(values=self.dance_moves, size=(30, 15), key='-MOVE_SELECT-', enable_events=True)],
            [sg.Text("Moves for selected lyric:"), sg.Text(size=(40, 1), key='-SELECTED_LYRIC_MOVES-')],
            [sg.Button("Submit Choreography", key='-SUBMIT_CHOREO-')]
        ]

    def save_tab(self):
        """Layout for the Save Choreography tab."""
        return [
            [sg.Text("Save your choreography to a Word file.")],
            [sg.Multiline(size=(50, 15), key='-PREVIEW-', disabled=True)],
            [sg.Button("Save to Word File", key='-SAVE-')]
        ]

    def add_dance_moves(self, moves_text):
        """Add multiple dance moves to the list."""
        moves = moves_text.strip().splitlines()
        for move in moves:
            move = move.strip()
            if move and move not in self.dance_moves:  # Avoid duplicates
                self.dance_moves.append(move)
        self.update_move_list()

    def update_move_list(self):
        """Update the list of dance moves."""
        self.window['-MOVE_LIST-'].update(values=self.dance_moves)
        self.window['-MOVE_SELECT-'].update(values=self.dance_moves)

    def submit_lyrics(self, lyrics_text):
        """Submit lyrics to be choreographed."""
        if lyrics_text.strip():
            self.lyrics = lyrics_text.strip().splitlines()
            # Initialize empty lists for moves for each lyric entry
            self.lyric_moves = [(lyric, []) for lyric in self.lyrics]
            self.update_lyric_list()

    def update_lyric_list(self):
        """Update the list of lyrics."""
        self.window['-LYRICS_LIST-'].update(values=self.lyrics)

    def assign_move_to_lyric(self, lyric_index):
        """Assign the selected move to the specified lyric."""
        if lyric_index is not None and self.current_selected_move:
            lyric, moves = self.lyric_moves[lyric_index]
            if len(moves) < 5:
                moves.append(self.current_selected_move)
                self.lyric_moves[lyric_index] = (lyric, moves)
                self.update_selected_lyric_moves(lyric_index)
            else:
                sg.popup("Limit Reached", "Each lyric can only have up to 5 moves.")

    def update_selected_lyric_moves(self, lyric_index):
        """Display the moves assigned to the selected lyric."""
        if lyric_index is not None:
            lyric, moves = self.lyric_moves[lyric_index]
            assigned_moves = ", ".join(moves)
            self.window['-SELECTED_LYRIC_MOVES-'].update(value=assigned_moves)

    def save_choreography(self, file_path):
        """Save the lyrics and associated moves to a Word file in table format."""
        if file_path:
            document = Document()
            document.add_heading("Choreography", level=1)

            # Add a table for lyrics and moves
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Lyric'
            hdr_cells[1].text = 'Dance Moves'

            # Add lyrics and moves to the table
            for lyric, moves in self.lyric_moves:
                row_cells = table.add_row().cells
                row_cells[0].text = lyric
                moves_text = ', '.join(moves) if moves else "No Move"
                row_cells[1].text = moves_text

            # Save the document
            document.save(file_path)
            sg.popup("Save Successful", f"Choreography saved to {file_path}")

    def toggle_move_selection(self, move):
        """Toggle the selection of a dance move."""
        if self.current_selected_move and move == self.current_selected_move:
            # Unselect the current move
            self.current_selected_move = None
        else:
            # Set the selected move
            self.current_selected_move = move

    def run(self):
        """Main loop to run the application."""
        while True:
            event, values = self.window.read()
            if event == sg.WINDOW_CLOSED:
                break

            elif event == '-ADD_MOVES-':
                moves_text = values['-MOVE_INPUT-']
                self.add_dance_moves(moves_text)
                self.window['-MOVE_INPUT-'].update('')

            elif event == '-SUBMIT_LYRICS-':
                lyrics_text = values['-LYRICS_INPUT-']
                self.submit_lyrics(lyrics_text)
                self.window['-LYRICS_INPUT-'].update('')

            elif event == '-MOVE_SELECT-':
                # Select a move to be assigned
                selected_move = values['-MOVE_SELECT-']
                if selected_move:
                    self.toggle_move_selection(selected_move[0])

            elif event == '-LYRICS_LIST-':
                # Retrieve the index of the selected lyric
                selected_lyric_index = values['-LYRICS_LIST-']
                if selected_lyric_index:
                    # We find the index based on the selected lyric
                    lyric_index = self.lyrics.index(selected_lyric_index[0])  # Get the actual index of the lyric
                    self.update_selected_lyric_moves(lyric_index)

            elif event == '-SUBMIT_CHOREO-':
                selected_lyric_index = values['-LYRICS_LIST-']
                if selected_lyric_index:
                    lyric_index = self.lyrics.index(selected_lyric_index[0])  # Get the index
                    if self.current_selected_move:
                        self.assign_move_to_lyric(lyric_index)

            elif event == '-SAVE-':
                file_path = sg.popup_get_file("Save As", save_as=True, no_window=True,
                                              file_types=(("Word files", "*.docx"),))
                if file_path:
                    self.save_choreography(file_path)

        self.window.close()


# Run the application
if __name__ == "__main__":
    app = DanceApp()
    app.run()
